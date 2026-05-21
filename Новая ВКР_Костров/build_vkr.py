"""Сборщик ВКР Кострова в формате .docx (Times New Roman 14, 1.5, поля 20-20-30-10 мм,
автособираемое оглавление, стили заголовков по методичке МФЮА).

Запуск: python3 build_vkr.py
Результат: ВКР_Костров_<YYYYMMDD_HHMM>.docx
"""
from __future__ import annotations

import datetime as _dt
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Mm, Pt, RGBColor

# Глобальные счётчик закладок и реестр заголовков (текст → имя закладки)
_BM_ID = [1000]
_BOOKMARKS: dict[str, str] = {}


def _next_bm_id() -> int:
    _BM_ID[0] += 1
    return _BM_ID[0]


def _bm_name(prefix: str) -> str:
    """Возвращает безопасное для Word имя закладки."""
    import re
    name = re.sub(r'[^A-Za-z0-9_]', '_', prefix)
    if not name or name[0].isdigit():
        name = '_' + name
    return name[:40]

from vkr_content import CONTENT  # длинный контент в отдельном файле для читаемости


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def configure_styles(doc: Document) -> None:
    # Normal
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), 'Times New Roman')
    pf = normal.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1 — главы и структурные элементы (ПРОПИСНЫЕ)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    rpr1 = h1.element.get_or_add_rPr()
    rfonts1 = rpr1.find(qn('w:rFonts'))
    if rfonts1 is None:
        rfonts1 = OxmlElement('w:rFonts')
        rpr1.append(rfonts1)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts1.set(qn(attr), 'Times New Roman')
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — подразделы (Sentence case)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    rpr2 = h2.element.get_or_add_rPr()
    rfonts2 = rpr2.find(qn('w:rFonts'))
    if rfonts2 is None:
        rfonts2 = OxmlElement('w:rFonts')
        rpr2.append(rfonts2)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts2.set(qn(attr), 'Times New Roman')
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Cm(1.25)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True


def set_page_geometry(doc: Document) -> None:
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        section.header_distance = Mm(12)
        section.footer_distance = Mm(12)


def add_page_numbers(doc: Document) -> None:
    """Добавляет нумерацию страниц по центру нижнего поля, без точки."""
    for section in doc.sections:
        footer = section.footer
        # очищаем
        for p in list(footer.paragraphs):
            p.clear()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE   \\* MERGEFORMAT'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r = run._element
        r.append(fldChar1)
        r.append(instrText)
        r.append(fldChar2)
        # Times New Roman 14
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def _wrap_with_bookmark(paragraph, bookmark_name: str) -> None:
    """Оборачивает содержимое параграфа парой bookmarkStart/bookmarkEnd."""
    bm_id = _next_bm_id()
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), bookmark_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    # вставляем bookmarkStart перед первым ребёнком, bookmarkEnd — после всего
    paragraph._element.insert(0, bm_start)
    paragraph._element.append(bm_end)


def add_heading1(doc: Document, text: str, *, all_caps: bool = True,
                 bookmark: str | None = None) -> None:
    """Заголовок 1 уровня — структурный элемент или глава."""
    page_break(doc)
    p = doc.add_paragraph(style='Heading 1')
    display = text.upper() if all_caps else text
    run = p.add_run(display)
    run.bold = True
    if bookmark:
        _wrap_with_bookmark(p, bookmark)
    # пустая строка после заголовка
    doc.add_paragraph()


def add_heading2(doc: Document, text: str, *, bookmark: str | None = None) -> None:
    p = doc.add_paragraph(style='Heading 2')
    p.add_run(text)
    if bookmark:
        _wrap_with_bookmark(p, bookmark)
    doc.add_paragraph()  # пустая строка после подзаголовка


def add_par(doc: Document, text: str, *, first_indent: bool = True,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    if not first_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(text)


def add_centered_par(doc: Document, text: str, *, bold: bool = False,
                     upper: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text.upper() if upper else text)
    run.bold = bold
    run.italic = italic


def add_blank(doc: Document, n: int = 1) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)


def _run_props(*, bold: bool = False) -> OxmlElement:
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), 'Times New Roman')
    rpr.append(rfonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '28')  # 14pt
    rpr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        b.set(qn('w:val'), 'true')
        rpr.append(b)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rpr.append(color)
    return rpr


def _add_toc_entry(doc: Document, text: str, bookmark: str, level: int) -> None:
    """Одна строка оглавления: гиперссылка → закладка + точечная заливка + PAGEREF."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if level == 2:
        p.paragraph_format.left_indent = Cm(0.75)
    # Правый табулятор с точечной заливкой на 16 см от левого поля
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark)
    hyperlink.set(qn('w:history'), '1')

    # 1) Текст заголовка
    r_text = OxmlElement('w:r')
    r_text.append(_run_props(bold=(level == 1)))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r_text.append(t)
    hyperlink.append(r_text)

    # 2) Табуляция (с точечной заливкой по tab_stop'у)
    r_tab = OxmlElement('w:r')
    r_tab.append(_run_props())
    tab = OxmlElement('w:tab')
    r_tab.append(tab)
    hyperlink.append(r_tab)

    # 3) Поле PAGEREF на закладку — реальный номер страницы
    r_fld_begin = OxmlElement('w:r')
    r_fld_begin.append(_run_props())
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    fc1.set(qn('w:dirty'), 'true')  # помечаем поле как «грязное» → пересчёт при открытии
    r_fld_begin.append(fc1)
    hyperlink.append(r_fld_begin)

    r_instr = OxmlElement('w:r')
    r_instr.append(_run_props())
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' PAGEREF {bookmark} \\h '
    r_instr.append(instr)
    hyperlink.append(r_instr)

    r_sep = OxmlElement('w:r')
    r_sep.append(_run_props())
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc2)
    hyperlink.append(r_sep)

    r_placeholder = OxmlElement('w:r')
    r_placeholder.append(_run_props())
    tn = OxmlElement('w:t')
    tn.text = '—'  # отображается до пересчёта поля
    r_placeholder.append(tn)
    hyperlink.append(r_placeholder)

    r_fld_end = OxmlElement('w:r')
    r_fld_end.append(_run_props())
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r_fld_end.append(fc3)
    hyperlink.append(r_fld_end)

    p._element.append(hyperlink)


def _set_update_fields_on_open(doc: Document) -> None:
    """Принудительная пересборка полей при открытии документа."""
    settings = doc.settings.element
    uf = OxmlElement('w:updateFields')
    uf.set(qn('w:val'), 'true')
    settings.append(uf)


def _collect_toc_entries() -> list[tuple[str, str, int]]:
    """(текст_в_оглавлении, имя_закладки, уровень)."""
    items: list[tuple[str, str, int]] = []
    items.append(('ВВЕДЕНИЕ', 'h_intro', 1))
    for ch_no, ch_title, sections in CONTENT['chapters']:
        items.append((f'ГЛАВА {ch_no}. {ch_title.upper()}', f'h_ch{ch_no}', 1))
        for sec_no, sec_title, _ in sections:
            items.append((f'{sec_no} {sec_title}',
                          f'h_sec_{sec_no.replace(".", "_")}', 2))
    items.append(('ЗАКЛЮЧЕНИЕ', 'h_concl', 1))
    items.append(('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', 'h_sources', 1))
    for i, (app_title, _) in enumerate(CONTENT['appendices'], 1):
        items.append((app_title.upper(), f'h_app_{i}', 1))
    return items


def add_table(doc: Document, header: list[str], rows: list[list[str]],
              caption: str, *, font_size: int = 12, line_spacing: float = 1.0) -> None:
    # подпись над таблицей
    cap = doc.add_paragraph()
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.line_spacing = 1.0
    r = cap.add_run(caption)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    # таблица
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = line_spacing
        run = p.add_run(h)
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
        run.bold = True
        _set_cell_borders(hdr_cells[i])
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = line_spacing
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
            _set_cell_borders(c)
    add_blank(doc, 1)


def add_formula(doc: Document, body: str, number: int, *,
                where: list[tuple[str, str]] | None = None,
                intro: str = '') -> None:
    """Оформление формулы по методичке п. 8.8."""
    if intro:
        add_par(doc, intro)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.add_run(body)
    # номер формулы в круглых скобках справа — добавим табуляцией
    pn = doc.add_paragraph()
    pn.paragraph_format.first_line_indent = Cm(0)
    pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pn.paragraph_format.line_spacing = 1.0
    pn.add_run(f'({number})')
    # пояснения
    if where:
        wp = doc.add_paragraph()
        wp.paragraph_format.first_line_indent = Cm(0)
        wp.paragraph_format.line_spacing = 1.0
        wp.add_run('где:')
        for sym, expl in where:
            ep = doc.add_paragraph()
            ep.paragraph_format.first_line_indent = Cm(0)
            ep.paragraph_format.line_spacing = 1.0
            ep.paragraph_format.left_indent = Cm(1.25)
            ep.add_run(f'{sym} — {expl};')


def build_title_page(doc: Document) -> None:
    add_centered_par(doc, 'АККРЕДИТОВАННОЕ ОБРАЗОВАТЕЛЬНОЕ ЧАСТНОЕ УЧРЕЖДЕНИЕ', bold=True)
    add_centered_par(doc, 'ВЫСШЕГО ОБРАЗОВАНИЯ', bold=True)
    add_centered_par(doc, '«МОСКОВСКИЙ ФИНАНСОВО-ЮРИДИЧЕСКИЙ УНИВЕРСИТЕТ МФЮА»', bold=True)
    add_blank(doc, 2)
    add_centered_par(doc, 'Кафедра менеджмента')
    add_centered_par(doc, 'Специальность 38.02.04 Коммерция (по отраслям)')
    add_blank(doc, 4)
    add_centered_par(doc, 'К ЗАЩИТЕ')
    add_centered_par(doc, '(РЕКОМЕНДОВАНО / НЕ РЕКОМЕНДОВАНО)')
    add_blank(doc, 1)
    add_centered_par(doc, 'Заведующий кафедрой')
    add_centered_par(doc, 'канд. экон. наук Е. П. Задворнева')
    add_centered_par(doc, '____________________ / Задворнева Е. П. /')
    add_centered_par(doc, '« ___ » ____________ 2026 г.')
    add_blank(doc, 3)
    add_centered_par(doc, 'ДИПЛОМНАЯ РАБОТА', bold=True, upper=True)
    add_blank(doc, 1)
    add_centered_par(doc, 'на тему:', italic=True)
    add_centered_par(doc, '«Перспективы применения интернет-технологий', bold=True)
    add_centered_par(doc, 'в коммерческой деятельности предприятий»', bold=True)
    add_blank(doc, 3)
    add_centered_par(doc, 'Обучающийся: Костров [И. О.]')
    add_centered_par(doc, '____________________ / Костров /')
    add_centered_par(doc, '« ___ » ____________ 2026 г.')
    add_blank(doc, 1)
    add_centered_par(doc, 'Индивидуальный номер обучающегося (ИНС): _________________')
    add_centered_par(doc, 'Группа: _________________')
    add_blank(doc, 2)
    add_centered_par(doc, 'Руководитель: канд. экон. наук [Ф. И. О.]')
    add_centered_par(doc, '____________________ / [Ф. И. О.] /')
    add_centered_par(doc, '« ___ » ____________ 2026 г.')
    add_blank(doc, 4)
    add_centered_par(doc, 'МОСКВА 2026', bold=True)


def build_task_page(doc: Document) -> None:
    page_break(doc)
    add_centered_par(doc, 'МОСКОВСКИЙ ФИНАНСОВО-ЮРИДИЧЕСКИЙ УНИВЕРСИТЕТ МФЮА', bold=True)
    add_centered_par(doc, 'Кафедра менеджмента')
    add_centered_par(doc, 'Направление / специальность 38.02.04 Коммерция (по отраслям)')
    add_blank(doc, 2)
    add_centered_par(doc, 'УТВЕРЖДАЮ')
    add_centered_par(doc, 'Заведующий кафедрой канд. экон. наук Е. П. Задворнева')
    add_centered_par(doc, '____________________')
    add_centered_par(doc, '« ___ » ____________ 2026 г.')
    add_blank(doc, 2)
    add_centered_par(doc, 'ЗАДАНИЕ НА ДИПЛОМНУЮ РАБОТУ', bold=True)
    add_blank(doc, 1)
    add_par(doc, 'Обучающийся: Костров [И. О.]')
    add_par(doc, 'Группа: _________________     Курс: 3')
    add_par(doc, 'Специальность: 38.02.04 Коммерция (по отраслям)')
    add_par(doc, 'Тема: «Перспективы применения интернет-технологий в коммерческой деятельности предприятий».')
    add_par(doc, 'Срок сдачи: « ___ » ____________ 2026 г.')
    add_par(doc, 'Место проведения преддипломной практики: Отдел экономической безопасности и противодействия коррупции УВД по ЮЗАО ГУ МВД России по г. Москве (ИНН организации 7727060703).')
    add_par(doc, 'Исходные данные: статистическая и аналитическая информация Росстата, Банка России, ФинЦЕРТ, Минцифры, АКИТ, Data Insight (2022–2025 гг.); открытые сведения о коммерческих предприятиях ЮЗАО; обезличенные материалы преддипломной практики.')
    add_par(doc, 'Оглавление расчётно-пояснительной записки:')
    add_par(doc, 'Введение. Глава 1. Теоретико-методологические основы применения интернет-технологий в коммерческой деятельности предприятий. Глава 2. Анализ применения интернет-технологий в коммерческой деятельности предприятий ЮЗАО (по материалам преддипломной практики в ОЭБиПК УВД по ЮЗАО). Глава 3. Перспективы применения интернет-технологий в коммерческой деятельности предприятий ЮЗАО и оценка эффективности предлагаемых концептуальных предложений. Заключение. Список использованных источников. Приложения.')
    add_par(doc, 'Рекомендуемая литература: учебно-методическая и нормативно-правовая литература, периодические издания по теме исследования (Сенчагов В. К., Авдийский В. И., Гончаренко Л. П., Шумилов В. М., Колесов Ю. И., Полякова Н. Б., Самарина Е. А., Александров О. Г. и др.), а также периодические издания «Финансы и кредит», «Российское предпринимательство», «Безопасность бизнеса», «Вестник МВД России», «Российский следователь».')
    add_blank(doc, 1)
    add_par(doc, 'Руководитель ДР: канд. экон. наук [Ф. И. О.]   ____________________   « ___ » ____________ 2026 г.')
    add_par(doc, 'Задание получил: Костров [И. О.]   ____________________   « ___ » ____________ 2026 г.')


def build_toc(doc: Document) -> None:
    # Заголовок «ОГЛАВЛЕНИЕ» — по центру, прописными, полужирно
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run('ОГЛАВЛЕНИЕ')
    run.bold = True

    # Статически собранные кликабельные строки оглавления
    for text, bookmark, level in _collect_toc_entries():
        _add_toc_entry(doc, text, bookmark, level)

    # Помечаем документ как требующий пересборки полей при открытии
    _set_update_fields_on_open(doc)


def build_introduction(doc: Document) -> None:
    add_heading1(doc, 'Введение', bookmark='h_intro')
    for par in CONTENT['intro']:
        add_par(doc, par)


def build_chapter(doc: Document, ch_no: int, ch_title: str, sections: list[tuple[str, str, list]]) -> None:
    add_heading1(doc, f'Глава {ch_no}. {ch_title}', bookmark=f'h_ch{ch_no}')
    for sec_no, sec_title, blocks in sections:
        add_heading2(doc, f'{sec_no} {sec_title}',
                     bookmark=f'h_sec_{sec_no.replace(".", "_")}')
        for block in blocks:
            kind = block[0]
            if kind == 'p':
                add_par(doc, block[1])
            elif kind == 'table':
                _, caption, header, rows = block
                add_table(doc, header, rows, caption)
            elif kind == 'formula':
                _, intro, body, number, where = block
                add_formula(doc, body, number, where=where, intro=intro)


def build_conclusion(doc: Document) -> None:
    add_heading1(doc, 'Заключение', bookmark='h_concl')
    for par in CONTENT['conclusion']:
        add_par(doc, par)


def build_sources(doc: Document) -> None:
    add_heading1(doc, 'Список использованных источников', bookmark='h_sources')
    for section_name, items in CONTENT['sources']:
        add_par(doc, section_name, first_indent=False)
        for i, src in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.add_run(f'{src}')


def build_appendices(doc: Document) -> None:
    for i, (app_title, paragraphs) in enumerate(CONTENT['appendices'], 1):
        add_heading1(doc, app_title, bookmark=f'h_app_{i}')
        for par in paragraphs:
            if isinstance(par, str):
                add_par(doc, par)
            elif isinstance(par, tuple) and par[0] == 'table':
                _, caption, header, rows = par
                add_table(doc, header, rows, caption)


def main() -> None:
    ts = _dt.datetime.utcnow().strftime('%Y%m%d_%H%M')
    out_dir = Path(__file__).parent
    out_name = f'ВКР_Костров_{ts}.docx'
    out_path = out_dir / out_name

    doc = Document()
    set_page_geometry(doc)
    configure_styles(doc)

    # По указанию пользователя титульный лист и задание на ДР в файле опущены —
    # эти страницы готовятся отдельно при сшивке работы (см. п. 11 методических
    # рекомендаций). Документ начинается сразу с оглавления.
    build_toc(doc)
    # 4. ВВЕДЕНИЕ
    build_introduction(doc)
    # 5. ГЛАВЫ
    for ch_no, ch_title, sections in CONTENT['chapters']:
        build_chapter(doc, ch_no, ch_title, sections)
    # 6. ЗАКЛЮЧЕНИЕ
    build_conclusion(doc)
    # 7. СПИСОК ИСТОЧНИКОВ
    build_sources(doc)
    # 8. ПРИЛОЖЕНИЯ
    build_appendices(doc)

    # Нумерация страниц
    add_page_numbers(doc)

    doc.save(out_path)
    print(f'Сохранено: {out_path}')


if __name__ == '__main__':
    main()
